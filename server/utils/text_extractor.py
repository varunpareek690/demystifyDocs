import io
import asyncio
from fastapi import UploadFile
from utils.exceptions import FileProcessingError

try:
    from pdfminer.high_level import extract_text as pdf_extract_text
    from pdfminer.layout import LAParams
except ImportError:
    try:
        import pypdf
    except ImportError:
        pdf_extract_text = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

class TextExtractor:
    """Utility class for extracting text from various file formats."""
    
    @staticmethod
    async def extract_text(file: UploadFile) -> str:
        """
        Extract text content from uploaded file.
        
        Args:
            file: FastAPI UploadFile object
            
        Returns:
            Extracted text content
            
        Raises:
            FileProcessingError: If text extraction fails
        """
        filename = file.filename or ""
        file_extension = TextExtractor._get_file_extension(filename).lower()
        
        try:
            # Read file content
            content = await file.read()
            
            # Extract text based on file type
            if file_extension == '.pdf':
                return await TextExtractor._extract_pdf_text(content)
            elif file_extension in ['.docx']:
                return await TextExtractor._extract_docx_text(content)
            elif file_extension in ['.doc']:
                # For .doc files, we'll attempt docx extraction (limited support)
                return await TextExtractor._extract_docx_text(content)
            elif file_extension == '.txt':
                return await TextExtractor._extract_txt_text(content)
            else:
                raise FileProcessingError(f"Unsupported file type: {file_extension}")
                
        except FileProcessingError:
            raise
        except Exception as e:
            raise FileProcessingError(f"Failed to extract text from file: {str(e)}")

    @staticmethod
    async def _extract_pdf_text(content: bytes) -> str:
        """Extract text from PDF content."""
        if pdf_extract_text is None:
            raise FileProcessingError("PDF processing library not available")
        
        try:
            # Use pdfminer for better text extraction
            def extract():
                return pdf_extract_text(
                    io.BytesIO(content),
                    laparams=LAParams(
                        all_texts=True,
                        detect_vertical=True,
                        word_margin=0.1,
                        char_margin=2.0,
                        line_margin=0.5,
                        boxes_flow=0.5
                    )
                )
            
            # Run in thread to avoid blocking
            text = await asyncio.get_event_loop().run_in_executor(None, extract)
            
            if not text.strip():
                # Fallback to pypdf if pdfminer returns empty
                try:
                    import pypdf
                    def pypdf_extract():
                        reader = pypdf.PdfReader(io.BytesIO(content))
                        text_parts = []
                        for page in reader.pages:
                            text_parts.append(page.extract_text())
                        return "\n".join(text_parts)
                    
                    text = await asyncio.get_event_loop().run_in_executor(None, pypdf_extract)
                except ImportError:
                    pass
            
            if not text.strip():
                raise FileProcessingError("No text content found in PDF")
                
            return text.strip()
            
        except Exception as e:
            raise FileProcessingError(f"Failed to extract text from PDF: {str(e)}")

    @staticmethod
    async def _extract_docx_text(content: bytes) -> str:
        """Extract text from DOCX/DOC content."""
        if DocxDocument is None:
            raise FileProcessingError("DOCX processing library not available")
        
        try:
            def extract():
                doc = DocxDocument(io.BytesIO(content))
                text_parts = []
                
                # Extract text from paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text.strip())
                
                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            text_parts.append(" | ".join(row_text))
                
                return "\n".join(text_parts)
            
            # Run in thread to avoid blocking
            text = await asyncio.get_event_loop().run_in_executor(None, extract)
            
            if not text.strip():
                raise FileProcessingError("No text content found in document")
                
            return text.strip()
            
        except Exception as e:
            raise FileProcessingError(f"Failed to extract text from document: {str(e)}")

    @staticmethod
    async def _extract_txt_text(content: bytes) -> str:
        """Extract text from TXT content."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    text = content.decode(encoding)
                    if text.strip():
                        return text.strip()
                except UnicodeDecodeError:
                    continue
            
            raise FileProcessingError("Unable to decode text file with any supported encoding")
            
        except Exception as e:
            raise FileProcessingError(f"Failed to extract text from file: {str(e)}")

    @staticmethod
    def _get_file_extension(filename: str) -> str:
        """Extract file extension from filename."""
        if '.' in filename:
            return '.' + filename.rsplit('.', 1)[1].lower()
        return ''

    @staticmethod
    def estimate_reading_time(text: str, words_per_minute: int = 200) -> int:
        """
        Estimate reading time for text in minutes.
        
        Args:
            text: Text content
            words_per_minute: Average reading speed
            
        Returns:
            Estimated reading time in minutes
        """
        word_count = len(text.split())
        return max(1, word_count // words_per_minute)

    @staticmethod
    def get_text_statistics(text: str) -> dict:
        """
        Get basic statistics about the text.
        
        Args:
            text: Text content
            
        Returns:
            Dict with text statistics
        """
        lines = text.split('\n')
        words = text.split()
        sentences = text.replace('!', '.').replace('?', '.').split('.')
        sentences = [s.strip() for s in sentences if s.strip()]
        
        return {
            'character_count': len(text),
            'word_count': len(words),
            'line_count': len(lines),
            'sentence_count': len(sentences),
            'average_words_per_sentence': len(words) / max(len(sentences), 1),
            'estimated_reading_time_minutes': TextExtractor.estimate_reading_time(text)
        }